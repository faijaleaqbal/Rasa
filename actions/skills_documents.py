import os
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
import requests

logger = logging.getLogger(__name__)

STORAGE_FILES_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "storage", "files"))
os.makedirs(STORAGE_FILES_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# 1. Telegram Document & Photo Dispatcher
# ---------------------------------------------------------------------------

def send_telegram_file(chat_id: str, file_path: str, caption: Optional[str] = None, file_type: str = "document") -> bool:
    """Sends a local file (document, photo, audio) to the user via Telegram Bot API."""
    bot_token = os.getenv("TELEGRAM_BOT_TOKEN")
    if not bot_token:
        logger.warning("No TELEGRAM_BOT_TOKEN found. Cannot send file to Telegram.")
        return False

    if not os.path.exists(file_path):
        logger.warning(f"File not found: {file_path}")
        return False

    try:
        if file_type == "photo" or file_path.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
            endpoint = f"https://api.telegram.org/bot{bot_token}/sendPhoto"
            field_name = "photo"
        elif file_type == "voice" or file_path.lower().endswith((".oga", ".ogg", ".opus")):
            endpoint = f"https://api.telegram.org/bot{bot_token}/sendVoice"
            field_name = "voice"
        else:
            endpoint = f"https://api.telegram.org/bot{bot_token}/sendDocument"
            field_name = "document"

        with open(file_path, "rb") as f:
            files = {field_name: (os.path.basename(file_path), f)}
            data = {"chat_id": str(chat_id)}
            if caption:
                try:
                    from addons.telegram_channel import format_telegram_markdown
                    norm_caption = format_telegram_markdown(caption)
                except Exception:
                    norm_caption = caption
                data["caption"] = norm_caption
                data["parse_mode"] = "Markdown"

            resp = requests.post(endpoint, data=data, files=files, timeout=30)
            if resp.status_code == 200:
                logger.info(f"Successfully sent {file_path} to Telegram chat {chat_id}")
                return True
            elif resp.status_code == 400 and caption:
                # If Telegram Markdown parse error, retry without parse_mode
                logger.warning(f"Telegram send file markdown failed ({resp.text}), retrying with plain text caption...")
                f.seek(0)
                data_plain = {"chat_id": str(chat_id), "caption": caption}
                resp_plain = requests.post(endpoint, data=data_plain, files={field_name: (os.path.basename(file_path), f)}, timeout=30)
                if resp_plain.status_code == 200:
                    return True
                logger.error(f"Telegram retry send file error: {resp_plain.text}")
            else:
                logger.error(f"Telegram send file error ({resp.status_code}): {resp.text}")
    except Exception as e:
        logger.error(f"Failed to send Telegram file: {e}")

    return False


# ---------------------------------------------------------------------------
# 2. PDF Handling (Read & Create)
# ---------------------------------------------------------------------------

def read_pdf_file(file_path: str, max_pages: int = 10) -> str:
    """Extracts text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        pages_to_read = min(total_pages, max_pages)

        extracted = []
        for i in range(pages_to_read):
            text = reader.pages[i].extract_text()
            if text:
                extracted.append(f"--- Page {i+1} ---\n{text.strip()}")

        if not extracted:
            return f"📄 PDF `{os.path.basename(file_path)}` has {total_pages} pages, but no readable text was extracted (might be scanned images)."

        summary = f"📄 **PDF: `{os.path.basename(file_path)}` ({pages_to_read}/{total_pages} pages extracted):**\n\n"
        summary += "\n\n".join(extracted)
        return summary[:3000] + ("\n... [Truncated for readability]" if len(summary) > 3000 else "")

    except Exception as e:
        logger.error(f"PDF read error: {e}")
        return f"❌ Failed to read PDF `{os.path.basename(file_path)}`: {str(e)}"


def create_pdf_file(title: str, content: str, filename: Optional[str] = None) -> Tuple[str, str]:
    """
    Creates a styled PDF document using ReportLab.
    Returns (file_path, message).
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.pdf"
        elif not filename.endswith(".pdf"):
            filename += ".pdf"

        file_path = os.path.join(STORAGE_FILES_DIR, filename)

        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Heading1"],
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#1A365D"),
            spaceAfter=12
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            textColor=colors.HexColor("#2D3748"),
            spaceAfter=10
        )
        date_style = ParagraphStyle(
            "DocDate",
            parent=styles["Italic"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#718096"),
            spaceAfter=15
        )

        story = [
            Paragraph(title, title_style),
            Paragraph(f"Generated by @Alya_Rasa_Bot • {datetime.now().strftime('%B %d, %Y %I:%M %p')}", date_style),
            HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E0"), spaceAfter=15)
        ]

        # Break content into paragraphs
        paragraphs = content.split("\n\n") if "\n\n" in content else content.split("\n")
        for p in paragraphs:
            clean_p = p.strip()
            if clean_p:
                story.append(Paragraph(clean_p.replace("\n", "<br/>"), body_style))
                story.append(Spacer(1, 6))

        doc.build(story)
        return file_path, f"✅ PDF document `{filename}` created successfully."

    except Exception as e:
        logger.error(f"PDF creation error: {e}")
        return "", f"❌ Failed to create PDF: {str(e)}"


# ---------------------------------------------------------------------------
# 3. Excel (.xlsx) Handling (Read & Create)
# ---------------------------------------------------------------------------

def read_excel_file(file_path: str, max_rows: int = 15) -> str:
    """Reads an Excel workbook and returns formatted tables."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheets_data = []

        for sheet_name in wb.sheetnames[:3]:
            ws = wb[sheet_name]
            rows_data = []
            for row in ws.iter_rows(values_only=True):
                if any(row):  # Not completely empty
                    rows_data.append([str(c) if c is not None else "" for c in row])
                if len(rows_data) >= max_rows:
                    break

            if rows_data:
                header = " | ".join(rows_data[0])
                divider = " | ".join(["---"] * len(rows_data[0]))
                table_rows = [" | ".join(r) for r in rows_data[1:]]
                sheet_md = f"📊 **Sheet: `{sheet_name}`**\n\n| {header} |\n| {divider} |\n"
                sheet_md += "\n".join([f"| {r} |" for r in table_rows])
                sheets_data.append(sheet_md)

        if not sheets_data:
            return f"📊 Excel file `{os.path.basename(file_path)}` contains no readable data."

        return f"📊 **Excel Data `{os.path.basename(file_path)}`:**\n\n" + "\n\n".join(sheets_data)

    except Exception as e:
        logger.error(f"Excel read error: {e}")
        return f"❌ Failed to read Excel file `{os.path.basename(file_path)}`: {str(e)}"


def create_excel_file(sheet_title: str, headers: List[str], rows: List[List[Any]], filename: Optional[str] = None) -> Tuple[str, str]:
    """
    Creates a styled Excel spreadsheet using openpyxl.
    Returns (file_path, message).
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"spreadsheet_{timestamp}.xlsx"
        elif not filename.endswith(".xlsx"):
            filename += ".xlsx"

        file_path = os.path.join(STORAGE_FILES_DIR, filename)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_title[:31]  # Sheet names capped at 31 chars

        # Header formatting
        header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

        ws.append(headers)
        for col_num in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align

        # Add rows
        row_font = Font(name="Calibri", size=10)
        thin_border = Border(
            left=Side(style="thin", color="E0E0E0"),
            right=Side(style="thin", color="E0E0E0"),
            top=Side(style="thin", color="E0E0E0"),
            bottom=Side(style="thin", color="E0E0E0")
        )

        for row_idx, r in enumerate(rows, start=2):
            ws.append(r)
            for col_idx in range(1, len(r) + 1):
                c = ws.cell(row=row_idx, column=col_idx)
                c.font = row_font
                c.border = thin_border
                if row_idx % 2 == 0:
                    c.fill = PatternFill(start_color="F2F4F8", end_color="F2F4F8", fill_type="solid")

        # Auto-adjust column widths
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

        wb.save(file_path)
        return file_path, f"✅ Excel spreadsheet `{filename}` created successfully with {len(rows)} rows."

    except Exception as e:
        logger.error(f"Excel creation error: {e}")
        return "", f"❌ Failed to create Excel file: {str(e)}"


# ---------------------------------------------------------------------------
# 4. Word (.docx) Handling (Read & Create)
# ---------------------------------------------------------------------------

def read_word_file(file_path: str) -> str:
    """Reads paragraphs and tables from a Word document using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
        text_parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                if p.style and p.style.name.startswith("Heading"):
                    text_parts.append(f"### {p.text.strip()}")
                else:
                    text_parts.append(p.text.strip())

        # Also parse tables
        for idx, table in enumerate(doc.tables, start=1):
            table_lines = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    table_lines.append(" | ".join(row_cells))
            if table_lines:
                text_parts.append(f"\n**[Table {idx}]**\n" + "\n".join(table_lines))

        if not text_parts:
            return f"📝 Word document `{os.path.basename(file_path)}` appears to be empty."

        content = "\n\n".join(text_parts)
        return f"📝 **Word Document `{os.path.basename(file_path)}`:**\n\n" + content[:3000]

    except Exception as e:
        logger.error(f"Word read error: {e}")
        return f"❌ Failed to read Word document `{os.path.basename(file_path)}`: {str(e)}"


def create_word_file(title: str, sections: List[Dict[str, str]], filename: Optional[str] = None) -> Tuple[str, str]:
    """
    Creates a styled Word (.docx) document using python-docx.
    `sections` is a list of dicts: [{"heading": "Section 1", "body": "Paragraph..."}]
    Returns (file_path, message).
    """
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.docx"
        elif not filename.endswith(".docx"):
            filename += ".docx"

        file_path = os.path.join(STORAGE_FILES_DIR, filename)

        doc = docx.Document()

        # Document Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_p.add_run(title)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(26, 54, 93)

        # Subtitle / Timestamp
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(f"Generated by @Alya_Rasa_Bot • {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
        sub_run.font.size = Pt(9)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(113, 128, 150)

        # Add sections
        for s in sections:
            h_text = s.get("heading")
            b_text = s.get("body", "")
            if h_text:
                doc.add_heading(h_text, level=2)
            if b_text:
                for line in b_text.split("\n"):
                    clean_line = line.strip()
                    if clean_line.startswith(("- ", "* ")):
                        p = doc.add_paragraph(clean_line[2:], style="List Bullet")
                    else:
                        p = doc.add_paragraph(clean_line)
                    p.paragraph_format.space_after = Pt(6)

        doc.save(file_path)
        return file_path, f"✅ Word document `{filename}` created successfully with {len(sections)} sections."

    except Exception as e:
        logger.error(f"Word creation error: {e}")
        return "", f"❌ Failed to create Word document: {str(e)}"
